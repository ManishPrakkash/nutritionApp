from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()

    # Title
    title = doc.add_heading('ZenFuel AI - Technical Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('Version: 1.0.0').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Tagline: Personalized Nutrition & Health Intelligence').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('---')

    # Section 1
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'ZenFuel AI is a premium, high-performance mobile application built with Flutter, '
        'designed to provide users with deep insights into their physiological state. '
        'It leverages machine learning models to predict health risks, calculate nutritional '
        'equilibrium, and generate adaptive meal and workout protocols.'
    )

    # Section 2
    doc.add_heading('2. Technology Stack', level=1)
    tech_stack = [
        ('Frontend Framework', 'Flutter (Dart SDK ^3.11.0)'),
        ('State Management', 'Riverpod (Unidirectional data flow, high testability)'),
        ('Database & Auth', 'Firebase (Cloud Firestore for NoSQL data, Firebase Auth for secure session management)'),
        ('Networking', 'Dio (Robust HTTP client for REST API communication)'),
        ('Typography', 'Google Fonts (Poppins)'),
        ('Data Visualization', 'fl_chart (Professional-grade health trend visualizations)')
    ]
    for item, Desc in tech_stack:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{item}: ')
        run.bold = True
        p.add_run(Desc)

    # Section 3
    doc.add_heading('3. Intelligent API Architecture', level=1)
    doc.add_paragraph('The application connects to a proprietary Python-based ML backend hosted on Render.')
    p = doc.add_paragraph()
    run = p.add_run('Base URL: ')
    run.bold = True
    p.add_run('https://zenhealth-app.onrender.com')

    doc.add_heading('Primary Endpoints:', level=2)
    
    endpoints = [
        ('1. /predict/all', 
         'Analyzes biometric data (Age, Gender, BMI, Activity Level). '
         'Uses Random Forest for BMI classification, Gradient Boosting for Badge status, '
         'and Ridge Regression for Daily Calorie Target calculation. '
         'Returns Percentage-based risk scores for Obesity, Diabetes, and Hypertension.'),
        ('2. /meal/plan', 
         'Generates breakfast, lunch, and dinner plans based on dietary preferences (Vegan, Keto, etc.), '
         'allergies, and medical history.'),
        ('3. /grocery/list', 
         'Compiles a categorized grocery list with filters for Periodicity, Budget, and People.'),
        ('4. /recipes/search', 
         'Reverse-searches recipes based on available household ingredients.')
    ]
    
    for title_text, Desc in endpoints:
        doc.add_heading(title_text, level=3)
        doc.add_paragraph(Desc)

    # Section 4
    doc.add_heading('4. Visual Components & Design System', level=1)
    doc.add_paragraph('The app follows a Luxury Light aesthetic, emphasizing clarity and precision.')
    
    doc.add_heading('Icons: Lucide Architecture', level=2)
    doc.add_paragraph(
        'The application uses a custom mapping system (lucide_fallback.dart) for Lucide Icon aesthetic '
        'coupled with Material Icons semantic bridges. Specific icons include flame (Caloric Burn), '
        'zap (BMR), shieldAlert (Health Risks), and cpu (AI Analysis).'
    )

    doc.add_heading('UI Components:', level=2)
    ui_comps = [
        'Glassmorphism Navigation Bar with BackdropFilter.',
        'Custom-Painted Biometric Gauges for BMI tracking.',
        'Neural Check-In momentum system for habits.'
    ]
    for comp in ui_comps:
        doc.add_paragraph(comp, style='List Bullet')

    # Section 5
    doc.add_heading('5. Data Flow & Security', level=1)
    security = [
        ('Local Persistence', 'Session data synchronized via Riverpod/Firestore.'),
        ('Session Isolation', 'All data tied to unique Firebase UID.'),
        ('Availability', 'Robust fallback logic for network disruptions.')
    ]
    for title_text, Desc in security:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(Desc)

    # Section 6
    doc.add_heading('6. Comprehensive Setup & Execution Guide', level=1)
    
    doc.add_heading('A. Environment Preparation', level=2)
    env_steps = [
        'Ensure Flutter SDK is installed and added to your system path.',
        'Verify setup by running: flutter doctor',
        'Initialize dependencies: flutter pub get'
    ]
    for step in env_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('B. Cloud & Authentication Setup', level=2)
    cloud_steps = [
        'Install Firebase CLI and authenticate.',
        'Configure the project: dart run flutterfire configure',
        'Ensure google-services.json is present in the android/app/ directory.'
    ]
    for step in cloud_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('C. Execution & Deployment', level=2)
    run_steps = [
        'Debug Mode: Connect a device/emulator and run "flutter run"',
        'Release Build: Execute "flutter build apk --release" for a distribution-ready package.',
        'Output Location: build/app/outputs/flutter-apk/app-release.apk'
    ]
    for step in run_steps:
        doc.add_paragraph(step, style='List Number')

    # Footer
    doc.add_paragraph('\n---')
    doc.add_paragraph('Generated by ZenFuel AI Technical Support').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save('ZenFuel_AI_Technical_Documentation.docx')
    print("Documentation saved successfully as ZenFuel_AI_Technical_Documentation.docx")

if __name__ == "__main__":
    create_docx()
